from docx import Document
from llama_cpp import Llama
import os

def get_model_path():
    directory = "LLM model"
    if not os.path.exists(directory):
        raise FileNotFoundError(f"The directory '{directory}' does not exist.")

    files = [f for f in os.listdir(directory) if os.path.isfile(os.path.join(directory, f))]

    if not files:
        raise FileNotFoundError(f"No files found in the '{directory}' directory.")
    if len(files) > 1:
        raise ValueError(f"Multiple files found in the '{directory}' directory. Ensure only one model file is present.")

    return os.path.join(directory, files[0])

# Initialize the model (Ensure the model path is correctly specified)
path = get_model_path()
model = Llama(path)


# Function to summarize a document
def summarize_document(file_path):
    # Extract text from the document
    document = Document(file_path)
    text = " ".join(para.text.strip() for para in document.paragraphs if para.text.strip())

    # final_prompt = f"Give an explanation of the following text in a lengthy structured paragraph. Go through every segment one by one, summarizing each. Include key points, attributes, and important features of the text, including article and section numbers and dates\n\n{text}"
    prompt = f"Summarize the following : {text}"
    try:
        # Generate a final summary from the combined summaries
        final_response = model.create_chat_completion(
            messages=[
                {"role": "system", "content": "You are a legal document summarizer assistant."},
                {"role": "user", "content": prompt}
            ]
        )
        return final_response['choices'][0]['message']['content'].strip()
    except Exception as e:
        return f"Error generating final summary: {str(e)}"

# Function to upload document and chat with it
def chat_with_upload_document(file_path):
    # Extract text from the document
    document = Document(file_path)
    text = " ".join(para.text.strip() for para in document.paragraphs if para.text.strip())

    prompt = f"i am giving you this information and only read this and dont give me output, just tell me if document is readed successfully or not, and i can ask questions related to the uploade document or not :\n\n{text}"

    try:
        # Generate a final summary from the combined summaries
        response = model.create_chat_completion(
            messages=[
                {"role": "system", "content": "You are a legal document reader assistant."},
                {"role": "user", "content": prompt}
            ]
        )
        return response['choices'][0]['message']['content'].strip()
    except Exception as e:
        return f"Error generating final summary: {str(e)}"

# Function to chat with the model
def chat_with_model(user_input):
    try:
        response = model.create_chat_completion(
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": user_input}
            ]
        )
        response = response['choices'][0]['message']['content']
        response = response.replace("<</SYS>>", "").strip()

    except Exception as e:
        response = f"Error in chat response: {str(e)}"

    return response

if __name__ == "__main__":
    print(chat_with_upload_document('Dummy_Regulation_Document.docx'))